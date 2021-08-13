from docx import Document
import docx
import random

lower= 'abcdefghijklmnopqrstwxyz'
upper=lower.upper()
numbers= '0123456789'
symbols= '_@!#$%^&*/?:'
alls= lower+upper+symbols+numbers
length=input("Enter length of password: ")
length=int(length)
service=input('Enter name of service: ')
service=service.upper()
password="".join(random.sample(alls,length))
print("Service: {}\n".format(service))
print("Password: {}".format(password))

doc = Document('/Users/swapnanildas/Desktop/Passwords.docx')

a= len(doc.paragraphs)

for para in doc.paragraphs:
  existing=para.text
  print(existing)

parat = [doc.paragraphs[i].text for i in range(a)]
doc = docx.Document()
for i in range(0,a):
    doc_para = doc.add_paragraph(parat[i])

doc_para = doc.add_heading('\nService: {}\nPassword: {}'.format(service,password),2)
doc.save('/Users/swapnanildas/Desktop/Passwords.docx')

